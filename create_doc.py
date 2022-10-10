import json
import os
import random
import time
import glob
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches
from PIL import Image

CONF = {}
# 解析配置文件
with open("subjectconf.json", "r", encoding="utf-8") as f:
# 获取配置字典
    CONF = json.load(f)
          
def handle_img(index,sub):
    """拼接图片"""
    src_path = './bank/'+ str(index) + '/'
    path_dir = os.listdir(src_path)
    filenum = len(path_dir) 
    if filenum > 0:
        picknum = sub["PICKNUM"]
        if filenum < picknum:
            picknum = filenum

        sample = random.sample(path_dir, picknum)
        imgs = []
        result = None

        for path in sample:
            if '.png' in path:
                des_path = src_path + path
                imgs.append(Image.open(des_path))
        
        if len(imgs)>0:
            # 单幅图像尺寸
            width, height = sub["UNIT_W"], sub["UNIT_H"]

            # 创建空白长图
            result = Image.new('RGB', (width*len(imgs), height), color=0)
            for i, img in enumerate(imgs):
                if img.mode == "P":
                    img = img.convert('RGB')

                img = img.resize((width, height))
                result.paste(img, box=(i*width,0))
    
        # 保存图片
        res_img = str(index)+'.png'

        if result:
            result.save(res_img)
        return res_img
    return ""

def create_doc():
    document = Document()
    date_now = datetime.now() 
    desk_path = os.path.join(os.path.expanduser("~"),'Desktop')+'/试卷/'
    
    if not os.path.exists(desk_path):
       os.makedirs(desk_path)
    res_doc = desk_path+'试卷'+ date_now.strftime('%Y-%m-%d %Hh%Mm%Ss')+'.docx'


    # 生成段落
    for index in CONF["subject_detail"]:
        # document.add_heading('Heading, level 1', level=1)
        document.add_paragraph(CONF["subject_detail"][index]["TEXT"], style='List Number')
        i=0
        while i < CONF["subject_detail"][index]["ROWS"]:
            img = handle_img(index, CONF["subject_detail"][index])
            if os.path.exists(img):
                document.add_picture(img,width = Inches(CONF["subject_detail"][index]["IMG_W"]),
                height=Inches(CONF["subject_detail"][index]["IMG_H"]))
                # document.add_picture(img,width = Inches(1),height=Inches(1))
            i+=1

    # 修改页边距
    document.sections[0].top_margin = Cm(2.5)
    document.sections[0].right_margin = Cm(1.27)
    document.sections[0].bottom_margin = Cm(2.5)
    document.sections[0].left_margin = Cm(1.27)

    document.save(res_doc)

def clear_cache():
    path = os.path.abspath(os.path.dirname(__file__))
    for cachefile in glob.glob(os.path.join(path, '*.png')):
        os.remove(cachefile)

def main():
    doc_num = CONF["doc_num"]
    for i in range(doc_num):
        create_doc()
        time.sleep(1)
    clear_cache()

if __name__ == '__main__':
    print("======开始生成试卷======")
    main()
    print("======试卷生成完毕======")
